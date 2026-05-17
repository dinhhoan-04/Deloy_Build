import io
import re
from datetime import datetime
from uuid import UUID

from docx import Document
from fastapi import APIRouter, Depends, Query, Response
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession

from app.deps import current_user, db
from app.repos.drafts import DraftRepo
from app.schemas.drafts import DraftCreate, DraftOut, DraftPatch
from rk_shared.models import User

router = APIRouter(prefix="/v1/drafts", tags=["drafts"])


def _out(d) -> DraftOut:
    return DraftOut(
        id=d.id,
        project_id=d.project_id,
        run_id=d.run_id,
        title=d.title,
        markdown=d.markdown,
        sections=d.sections,
        created_at=d.created_at,
        updated_at=d.updated_at,
    )


def _build_docx(title: str, markdown: str, updated_at: datetime) -> bytes:
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Last updated: {updated_at.strftime('%Y-%m-%d')}")
    for line in markdown.splitlines():
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.strip():
            doc.add_paragraph(line.strip())
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.post("", response_model=DraftOut, status_code=201)
async def upsert_draft(
    body: DraftCreate,
    u: User = Depends(current_user),
    s: AsyncSession = Depends(db),
):
    repo = DraftRepo(s)
    draft = await repo.upsert(
        u.id,
        project_id=body.project_id,
        run_id=body.run_id,
        title=body.title,
        markdown=body.markdown,
        sections=body.sections,
    )
    await s.commit()
    return _out(draft)


@router.get("", response_model=DraftOut)
async def get_draft(
    project_id: UUID = Query(...),
    u: User = Depends(current_user),
    s: AsyncSession = Depends(db),
):
    return _out(await DraftRepo(s).get(u.id, project_id))


@router.patch("/{draft_id}", response_model=DraftOut)
async def patch_draft(
    draft_id: UUID,
    body: DraftPatch,
    u: User = Depends(current_user),
    s: AsyncSession = Depends(db),
):
    draft = await DraftRepo(s).patch(u.id, draft_id, title=body.title, markdown=body.markdown)
    await s.commit()
    return _out(draft)


@router.delete("/{draft_id}", status_code=204)
async def delete_draft(
    draft_id: UUID,
    u: User = Depends(current_user),
    s: AsyncSession = Depends(db),
) -> Response:
    await DraftRepo(s).delete(u.id, draft_id)
    await s.commit()
    return Response(status_code=204)


@router.get("/{draft_id}/export")
async def export_draft(
    draft_id: UUID,
    format: str = Query(default="md", pattern="^(md|docx)$"),
    u: User = Depends(current_user),
    s: AsyncSession = Depends(db),
):
    draft = await DraftRepo(s).get_by_id(u.id, draft_id)
    safe_title = re.sub(r"[^\w\s-]", "", draft.title).strip().replace(" ", "_") or "draft"
    date_str = draft.updated_at.strftime("%Y-%m-%d")

    if format == "docx":
        data = _build_docx(draft.title, draft.markdown, draft.updated_at)
        return StreamingResponse(
            io.BytesIO(data),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'},
        )

    frontmatter = f"---\ntitle: {draft.title}\ndate: {date_str}\n---\n\n"
    return Response(
        content=frontmatter + draft.markdown,
        media_type="text/markdown",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.md"'},
    )
