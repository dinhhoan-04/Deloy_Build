import datetime
import io
import uuid
from app.schemas.drafts import DraftCreate, DraftPatch, DraftOut


def test_draft_create_requires_markdown():
    d = DraftCreate(
        project_id=uuid.uuid4(),
        markdown="# Hello",
    )
    assert d.title == "Untitled Draft"
    assert d.sections == []


def test_draft_patch_all_optional():
    p = DraftPatch()
    assert p.title is None
    assert p.markdown is None


def test_draft_out_round_trip():
    import datetime
    now = datetime.datetime.utcnow()
    d = DraftOut(
        id=uuid.uuid4(),
        project_id=uuid.uuid4(),
        run_id=None,
        title="My Draft",
        markdown="# Hello",
        sections=[],
        created_at=now,
        updated_at=now,
    )
    assert d.title == "My Draft"


def test_build_docx_includes_date():
    from docx import Document
    from app.routers.drafts import _build_docx
    updated = datetime.datetime(2026, 5, 15, 12, 0, 0)
    data = _build_docx("My Title", "Some content", updated)
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs]
    assert any("2026-05-15" in p for p in paragraphs)


def test_build_docx_title_present():
    from docx import Document
    from app.routers.drafts import _build_docx
    updated = datetime.datetime(2026, 5, 15)
    data = _build_docx("Test Draft", "# Heading\n\nBody text", updated)
    doc = Document(io.BytesIO(data))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "Test Draft" in full_text
