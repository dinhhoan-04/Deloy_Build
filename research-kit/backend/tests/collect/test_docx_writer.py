import io
from docx import Document
from app.schemas import MergedClaim, CollectCitation
from app.services.collect.docx_writer import build_docx_bytes


def _claim(text, score=None):
    return MergedClaim(
        claim_text=text,
        citations=[CollectCitation(ref_id="1", url="https://doi.org/A")],
        source_tools=["elicit"],
        verify_score=score,
    )


def test_build_docx_returns_valid_docx_with_topic_heading():
    data = build_docx_bytes(
        topic="Sleep", date="2026-04-30", tools_used=["elicit"], run_verify=False,
        summary="S.", claims=[_claim("hi")], references=[], raw_by_tool={"elicit": "raw"},
    )
    doc = Document(io.BytesIO(data))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("Sleep" in h for h in headings)


def test_build_docx_groups_by_score_when_verify_active():
    data = build_docx_bytes(
        topic="x", date="d", tools_used=["elicit"], run_verify=True,
        summary="s", claims=[_claim("high", 90), _claim("low", 20)], references=[], raw_by_tool={},
    )
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "High-confidence" in text and "Low-confidence" in text
